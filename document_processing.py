from __future__ import annotations

import io
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import fitz  # PyMuPDF
import docx
import pandas as pd


@dataclass
class RawDocument:
    source: str
    page: str
    text: str


@dataclass
class Chunk:
    chunk_id: str
    source: str
    page: str
    text: str


def clean_text(text: str) -> str:
    text = re.sub(r"\s+", " ", text or "").strip()
    return text


def read_pdf(file_bytes: bytes, filename: str) -> list[RawDocument]:
    docs: list[RawDocument] = []
    pdf = fitz.open(stream=file_bytes, filetype="pdf")
    for i, page in enumerate(pdf, start=1):
        text = clean_text(page.get_text("text"))
        if text:
            docs.append(RawDocument(source=filename, page=str(i), text=text))
    return docs


def read_docx(file_bytes: bytes, filename: str) -> list[RawDocument]:
    f = io.BytesIO(file_bytes)
    d = docx.Document(f)
    paras = [p.text for p in d.paragraphs if p.text.strip()]
    text = clean_text("\n".join(paras))
    return [RawDocument(source=filename, page="Not available", text=text)] if text else []


def read_txt(file_bytes: bytes, filename: str) -> list[RawDocument]:
    text = file_bytes.decode("utf-8", errors="ignore")
    text = clean_text(text)
    return [RawDocument(source=filename, page="Not available", text=text)] if text else []


def read_xlsx(file_bytes: bytes, filename: str) -> list[RawDocument]:
    f = io.BytesIO(file_bytes)
    xl = pd.ExcelFile(f)
    docs: list[RawDocument] = []
    for sheet in xl.sheet_names:
        df = xl.parse(sheet).fillna("")
        rows = []
        for _, row in df.iterrows():
            row_text = " | ".join(str(x) for x in row.tolist() if str(x).strip())
            if row_text.strip():
                rows.append(row_text)
        text = clean_text("\n".join(rows))
        if text:
            docs.append(RawDocument(source=f"{filename} / Sheet: {sheet}", page="Not available", text=text))
    return docs


def read_uploaded_file(file) -> list[RawDocument]:
    filename = file.name
    suffix = Path(filename).suffix.lower()
    file_bytes = file.getvalue()
    if suffix == ".pdf":
        return read_pdf(file_bytes, filename)
    if suffix == ".docx":
        return read_docx(file_bytes, filename)
    if suffix in {".txt", ".md", ".csv"}:
        return read_txt(file_bytes, filename)
    if suffix in {".xlsx", ".xls"}:
        return read_xlsx(file_bytes, filename)
    return read_txt(file_bytes, filename)


def chunk_documents(raw_docs: Iterable[RawDocument], chunk_words: int = 260, overlap_words: int = 50) -> list[Chunk]:
    chunks: list[Chunk] = []
    counter = 1
    for doc in raw_docs:
        words = doc.text.split()
        if not words:
            continue
        step = max(1, chunk_words - overlap_words)
        for start in range(0, len(words), step):
            piece = words[start:start + chunk_words]
            if len(piece) < 30:
                continue
            chunks.append(
                Chunk(
                    chunk_id=f"C{counter:03d}",
                    source=doc.source,
                    page=doc.page,
                    text=" ".join(piece),
                )
            )
            counter += 1
    return chunks
